import tkinter as tk
from tkinter import ttk, filedialog
from docx import Document


def create_education_fields(num_education)
    global edu_degree_entries, edu_major_entries, edu_institution_entries, edu_years_entries
    edu_degree_entries = []
    edu_major_entries = []
    edu_institution_entries = []
    edu_years_entries = []

    for i in range(num_education)
        row_offset = 13 + i  4  # Adjusting row based on number of fields added
        tk.Label(app, text=fEducation {i + 1} Degree).grid(row=row_offset, column=0, sticky=e, padx=(5, 0))
        degree_entry = tk.Entry(app)
        degree_entry.grid(row=row_offset, column=1)
        edu_degree_entries.append(degree_entry)

        tk.Label(app, text=fEducation {i + 1} Major).grid(row=row_offset + 1, column=0, sticky=e, padx=(5, 0))
        major_entry = tk.Entry(app)
        major_entry.grid(row=row_offset + 1, column=1)
        edu_major_entries.append(major_entry)

        tk.Label(app, text=fEducation {i + 1} Institution).grid(row=row_offset + 2, column=0, sticky=e, padx=(5, 0))
        institution_entry = tk.Entry(app)
        institution_entry.grid(row=row_offset + 2, column=1)
        edu_institution_entries.append(institution_entry)

        tk.Label(app, text=fEducation {i + 1} Years).grid(row=row_offset + 3, column=0, sticky=e, padx=(5, 0))
        years_entry = tk.Entry(app)
        years_entry.grid(row=row_offset + 3, column=1)
        edu_years_entries.append(years_entry)


def create_work_experience_fields(num_experiences)
    global work_role_entries, work_company_entries, work_years_entries
    work_role_entries = []
    work_company_entries = []
    work_years_entries = []

    for i in range(num_experiences)
        row_offset = 13 + num_education  4 + i  3  # Adjusting row based on education fields
        tk.Label(app, text=fWork Experience {i + 1} Role).grid(row=row_offset, column=0, sticky=e, padx=(5, 0))
        role_entry = tk.Entry(app)
        role_entry.grid(row=row_offset, column=1)
        work_role_entries.append(role_entry)

        tk.Label(app, text=fWork Experience {i + 1} Company).grid(row=row_offset + 1, column=0, sticky=e, padx=(5, 0))
        company_entry = tk.Entry(app)
        company_entry.grid(row=row_offset + 1, column=1)
        work_company_entries.append(company_entry)

        tk.Label(app, text=fWork Experience {i + 1} Years).grid(row=row_offset + 2, column=0, sticky=e, padx=(5, 0))
        years_entry = tk.Entry(app)
        years_entry.grid(row=row_offset + 2, column=1)
        work_years_entries.append(years_entry)


def adjust_alignment()
    for widget in app.grid_slaves()
        widget.grid_configure(sticky=w)


def generate_fields()
    global num_education, num_experiences
    num_education = int(num_education_entry.get())
    num_experiences = int(num_experiences_entry.get())
    create_education_fields(num_education)
    create_work_experience_fields(num_experiences)
    skills_label.grid(row=35 + num_education  4 + num_experiences  3, column=0, sticky=w, pady=5)
    skills_entry.grid(row=35 + num_education  4 + num_experiences  3 + 1, column=0, columnspan=2)
    references_label.grid(row=37 + num_education  4 + num_experiences  3, column=0, sticky=w, pady=5)
    references_entry.grid(row=37 + num_education  4 + num_experiences  3 + 1, column=0, columnspan=2)
    submit_button.grid(row=39 + num_education  4 + num_experiences  3, column=0, columnspan=2)
    adjust_alignment()


def submit_details()
    # Create a Word document for the resume
    document = Document()

    # Add Name
    document.add_heading(name_entry.get(), level=1)

    # Add Contact Information
    contact_info = f{location_entry.get()}  {email_entry.get()}  {linkedin_entry.get()}  {portfolio_entry.get()}
    document.add_paragraph(contact_info, style=Normal)

    # Add Summary
    document.add_heading(Summary, level=2)
    document.add_paragraph(summary_entry.get(1.0, end).strip())

    # Add Work Experience
    document.add_heading(Work Experience, level=2)
    for i in range(num_experiences)
        role = work_role_entries[i].get()
        company = work_company_entries[i].get()
        years = work_years_entries[i].get()
        document.add_paragraph(f{role} at {company} ({years}), style=Normal)
        # Add Education
        document.add_heading(Education, level=2)
        for i in range(num_education)
            degree = edu_degree_entries[i].get()
            major = edu_major_entries[i].get()
            institution = edu_institution_entries[i].get()
            years = edu_years_entries[i].get()
            document.add_paragraph(f{degree} in {major}, {institution} ({years}), style=Normal)

        # Add Skills
        document.add_heading(Skills, level=2)
        skills = skills_entry.get().split(,)  # Assume skills are comma-separated
        for skill in skills
            document.add_paragraph(f• {skill.strip()}, style=List Bullet)

        # Add References
        document.add_heading(References, level=2)
        document.add_paragraph(references_entry.get(), style=Normal)

        # Save the document
        file_path = filedialog.asksaveasfilename(defaultextension=.docx,
                                                 filetypes=[(Word Documents, .docx)],
                                                 title=Save Resume As)
        if file_path
            document.save(file_path)


# Initialize the tkinter app
app = tk.Tk()
app.title(Resume Maker)

# Personal Information
tk.Label(app, text=Name, font=(Helvetica, 14, bold)).grid(row=0, column=0, sticky=w, padx=5, pady=5)
name_entry = tk.Entry(app, font=(Helvetica, 14))
name_entry.grid(row=0, column=1, padx=5, pady=5)

tk.Label(app, text=Location (City, Country)).grid(row=1, column=0, sticky=e, padx=5)
location_entry = tk.Entry(app)
location_entry.grid(row=1, column=1)

tk.Label(app, text=Email Address).grid(row=2, column=0, sticky=e, padx=5)
email_entry = tk.Entry(app)
email_entry.grid(row=2, column=1)

tk.Label(app, text=LinkedIn Profile).grid(row=3, column=0, sticky=e, padx=5)
linkedin_entry = tk.Entry(app)
linkedin_entry.grid(row=3, column=1)

tk.Label(app, text=Portfolio Link).grid(row=4, column=0, sticky=e, padx=5)
portfolio_entry = tk.Entry(app)
portfolio_entry.grid(row=4, column=1)

# Summary Section
tk.Label(app, text=Summary, font=(Helvetica, 12, bold)).grid(row=5, column=0, sticky=w, pady=5)
summary_entry = tk.Text(app, height=4, width=50)
summary_entry.grid(row=6, column=0, columnspan=2)

# Number of Entries
tk.Label(app, text=Number of Education Entries).grid(row=7, column=0, sticky=e, padx=5)
num_education_entry = tk.Entry(app)
num_education_entry.grid(row=7, column=1)

tk.Label(app, text=Number of Work Experience Entries).grid(row=8, column=0, sticky=e, padx=5)
num_experiences_entry = tk.Entry(app)
num_experiences_entry.grid(row=8, column=1)

# Generate Fields Button
generate_button = tk.Button(app, text=Generate Fields, command=generate_fields)
generate_button.grid(row=9, column=0, columnspan=2)

# Skills Section
skills_label = tk.Label(app, text=Skills, font=(Helvetica, 12, bold))
skills_entry = tk.Entry(app)

# References Section
references_label = tk.Label(app, text=References, font=(Helvetica, 12, bold))
references_entry = tk.Entry(app)

# Submit Button
submit_button = tk.Button(app, text=Submit, command=submit_details)

# Start the application
app.mainloop()
